# 合并文档
import docx
import Path

world_files_path = ""


def merge_doc(docx_files: list):
    for docx_file in sorted(docx_files):
        anthor_doc = Document(docx_file)
        paras = anthor_doc.paragraphs
        # paras_content = [paras.text for para in paras]
        for para in paras:
            newpara = doc.add_paragraph('')
            newpara.add_run(para.text)

    doc.save(Path(world_files_path, 'new.docx'))


if __name__ == 'main':
    print("main call")
    files = ""
    merge_doc(files)